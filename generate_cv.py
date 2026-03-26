#!/usr/bin/env python3
"""Generate a professional ATS-friendly 2-page Word CV for Hamid Raza."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page Setup (tighter margins for 2-page fit) ──
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.6)
    section.right_margin = Cm(1.6)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(1)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.08

# ── Colors (Blue theme) ──
BLUE = RGBColor(0x3B, 0x82, 0xF6)
DARK = RGBColor(0x0F, 0x17, 0x2A)
GRAY = RGBColor(0x64, 0x74, 0x8B)
TEXT = RGBColor(0x33, 0x41, 0x55)


def add_hyperlink(paragraph, text, url, color=BLUE, size=Pt(9)):
    """Add a clickable hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} r:id="{r_id}" xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"></w:hyperlink>')
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'<w:rPr><w:rFonts w:ascii="Calibri" w:hAnsi="Calibri"/><w:sz w:val="{int(size.pt * 2)}"/><w:color w:val="{str(color)}"/><w:u w:val="single"/></w:rPr>'
        f'<w:t xml:space="preserve">{text}</w:t>'
        f'</w:r>'
    )
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_line(doc, color=BLUE, sz=2):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="{sz*4}" w:space="1" w:color="{str(color)}"/>'
        f'</w:pBdr>'
    ))


def heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text.upper())
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'
    add_line(doc, BLUE, 1)


def experience(doc, role, company, date, location, bullets):
    # Role + Date on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(role)
    run.font.size = Pt(10.5)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(2)
    run = p2.add_run(company)
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run.font.name = 'Calibri'
    run = p2.add_run(f"  |  {date}  |  {location}")
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'

    for bullet in bullets:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Cm(0.6)
        run = p.add_run(bullet)
        run.font.size = Pt(9.5)
        run.font.color.rgb = TEXT
        run.font.name = 'Calibri'


def project(doc, name, desc, tech):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(name)
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = DARK
    run.font.name = 'Calibri'
    run = p.add_run(f"  [{tech}]")
    run.font.size = Pt(8.5)
    run.font.color.rgb = GRAY
    run.font.name = 'Calibri'
    run = p.add_run(f" - {desc}")
    run.font.size = Pt(9)
    run.font.color.rgb = TEXT
    run.font.name = 'Calibri'


# ══════════════════════════════════════════
# HEADER
# ══════════════════════════════════════════

# Header table: Name/Title/Contact on left, Photo on right
header_table = doc.add_table(rows=1, cols=2)
header_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Info cell (left)
info_cell = header_table.rows[0].cells[0]
info_cell.paragraphs[0].clear()

p = info_cell.paragraphs[0]
p.paragraph_format.space_after = Pt(1)
run = p.add_run("HAMID RAZA")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = DARK
run.font.name = 'Calibri'

p = info_cell.add_paragraph()
p.paragraph_format.space_after = Pt(3)
run = p.add_run("Ruby on Rails - Full Stack Developer")
run.font.size = Pt(12)
run.font.color.rgb = BLUE
run.font.name = 'Calibri'
run.font.bold = True

p = info_cell.add_paragraph()
p.paragraph_format.space_after = Pt(1)
run = p.add_run("+92 322 4513443  |  mr.hamid.raza@gmail.com")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

p = info_cell.add_paragraph()
p.paragraph_format.space_after = Pt(0)
add_hyperlink(p, "linkedin.com/in/hamidraza84-rubyonrails", "https://www.linkedin.com/in/hamidraza84-rubyonrails/", BLUE, Pt(9))
run = p.add_run("  |  Lahore, Pakistan")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

# Photo cell (right - top aligned)
photo_cell = header_table.rows[0].cells[1]
photo_cell.width = Cm(3.5)
# Top-align the photo cell
tc = photo_cell._tc
tcPr = tc.get_or_add_tcPr()
tcPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="top"/>'))
photo_p = photo_cell.paragraphs[0]
photo_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
photo_p.paragraph_format.space_before = Pt(0)
photo_p.paragraph_format.space_after = Pt(0)
script_dir = os.path.dirname(os.path.abspath(__file__))
photo_path = os.path.join(script_dir, "profile.jpg")
if os.path.exists(photo_path):
    from PIL import Image, ImageDraw, ImageOps
    import io

    # Create circular cropped image with border
    img = Image.open(photo_path).convert("RGBA")
    # Crop to square
    size = min(img.size)
    left = (img.width - size) // 2
    top = (img.height - size) // 2
    img = img.crop((left, top, left + size, top + size))
    img = img.resize((500, 500), Image.LANCZOS)

    # Create circular mask
    mask = Image.new("L", (500, 500), 0)
    draw = ImageDraw.Draw(mask)
    draw.ellipse((0, 0, 500, 500), fill=255)

    # Add border ring (ruby red)
    border_width = 12
    border_img = Image.new("RGBA", (500, 500), (0, 0, 0, 0))
    border_draw = ImageDraw.Draw(border_img)
    border_draw.ellipse((0, 0, 499, 499), outline=(59, 130, 246, 255), width=border_width)

    # Apply mask to photo
    output = Image.new("RGBA", (500, 500), (255, 255, 255, 0))
    output.paste(img, (0, 0), mask)
    output.paste(border_img, (0, 0), border_img)

    # Save to bytes
    img_bytes = io.BytesIO()
    output.save(img_bytes, format="PNG")
    img_bytes.seek(0)

    photo_p.add_run().add_picture(img_bytes, width=Cm(3))

# Remove borders from header table
for cell in header_table.rows[0].cells:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcPr.append(parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tcBorders>'
    ))

add_line(doc, BLUE, 2)

# ══════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ══════════════════════════════════════════

heading(doc, "Professional Summary")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
run = p.add_run(
    "Seasoned Full Stack Engineer with 10+ years of experience in Ruby on Rails and 2+ years building backend services "
    "with NestJS and TypeScript. Skilled in designing, developing, and optimizing scalable healthcare and eCommerce web applications. "
    "Proficient in building secure RESTful and GraphQL APIs, integrating healthcare data standards (FHIR/HL7), "
    "and ensuring HIPAA compliance. Proven track record delivering 17+ projects across 6+ countries for clients "
    "in eCommerce, healthcare, logistics, IoT, and SaaS domains."
)
run.font.size = Pt(9.5)
run.font.color.rgb = TEXT
run.font.name = 'Calibri'

# ══════════════════════════════════════════
# CORE COMPETENCIES
# ══════════════════════════════════════════

heading(doc, "Core Competencies")

table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

skills = [
    ("Backend: Ruby on Rails, NestJS, TypeScript, Solidus/Spree, RESTful APIs, GraphQL, TypeORM",
     "Frontend: React.js, Vue.js, Hotwire/Turbo, ActionCable, Bootstrap, SASS/SCSS"),
    ("Databases: PostgreSQL, MySQL, Redis, Elasticsearch, N+1 Elimination, Query Optimization",
     "Background: Sidekiq, ActiveJob, RabbitMQ, Redis Caching, Fragment Caching"),
    ("DevOps: AWS, Docker, Heroku, CI/CD, GitHub Actions, Jenkins",
     "Payments & Healthcare: Stripe/Connect, PayPal, Adyen, FHIR/HL7, HIPAA, RBAC"),
]

for i, (left, right) in enumerate(skills):
    for j, text in enumerate([left, right]):
        cell = table.rows[i].cells[j]
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(1)
        p.runs  # init
        run = p.add_run("\u2022  " + text)
        run.font.size = Pt(8.5)
        run.font.color.rgb = TEXT
        run.font.name = 'Calibri'
        # Remove borders
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcPr.append(parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        ))

# ══════════════════════════════════════════
# PROFESSIONAL EXPERIENCE
# ══════════════════════════════════════════

heading(doc, "Professional Experience")

experience(doc,
    "Senior Backend Developer",
    "Valus.io",
    "2024 - Present", "Saudi Arabia (Remote)",
    [
        "Architecting scalable Rails 7 APIs for a large-scale eCommerce grocery platform serving millions of users",
        "Hands-on ownership of Solidus/Spree commerce domains: orders, inventory, promotions, pricing, taxes, shipments, returns",
        "Implemented real-time order tracking and notifications using ActionCable, backed by Sidekiq background jobs",
        "Performance tuning across large datasets using N+1 elimination, Redis caching, and pagination strategies",
        "AWS-based deployments and CI/CD pipelines with focus on reliability, observability, and zero-downtime releases",
    ]
)

experience(doc,
    "NestJS Backend Developer",
    "Freelance",
    "Jan 2024 - Present", "Remote",
    [
        "Built RESTful API backends for client projects using NestJS, TypeScript, and TypeORM with PostgreSQL",
        "Developed CAST4k Reseller Panel backend with Stripe payment integration, subscription management, and license key generation",
        "Created a personal finance tracking application with budget management, transaction categorization, and reporting",
        "Implemented JWT authentication, role-based access control, input validation, and Swagger API documentation",
    ]
)

experience(doc,
    "Senior Backend Developer",
    "Tractive",
    "2022 - 2024", "Austria (Remote)",
    [
        "Developed backend infrastructure for a real-time GPS pet tracking system processing millions of location events",
        "Designed Rails APIs for location and event data with sub-100ms response times for mobile clients",
        "Optimized database performance and reduced hot-path overhead through Redis caching and query optimization",
        "Contributed to event-driven workflows integrating RabbitMQ with Rails services for async processing",
    ]
)

experience(doc,
    "Full Stack Developer",
    "Doc.gr",
    "2022 - 2023", "Greece (Remote)",
    [
        "Built healthcare-focused appointment scheduling and doctor-patient management platform using Rails and React.js",
        "Implemented HIPAA-compliant workflows, secure patient data handling, and audit trail processes",
        "Integrated FHIR/HL7 healthcare data exchange standards through Rails API layers for interoperability",
    ]
)

experience(doc,
    "Senior Backend Developer",
    "SpaceOS",
    "2021 - 2022", "Poland (Remote)",
    [
        "Developed intelligent workplace automation platform with IoT integration and multi-tenant SaaS architecture",
        "Implemented real-time notifications using ActionCable and Turbo Streams for instant workspace updates",
        "Designed tenant isolation patterns and RBAC for secure multi-organization data management",
    ]
)

experience(doc,
    "Full Stack Developer",
    "Borderhaul",
    "2021 - 2022", "USA (Remote)",
    [
        "Built logistics and freight platform connecting shippers with cross-border carriers across North America",
        "Automated payouts via Stripe Connect; built real-time shipment tracking with ActionCable and webhooks",
        "Developed quote engine, carrier matching algorithms, and automated billing workflows",
    ]
)

experience(doc,
    "Senior Full Stack Developer",
    "Coeus Solutions",
    "2018 - 2021", "Lahore, Pakistan",
    [
        "Led Rails SaaS development end-to-end: data modeling, API design, background jobs, production hardening",
        "Built scalable REST APIs, multi-currency payment gateways (Stripe, PayPal, Adyen), and Hotwire/Turbo UI",
        "Introduced performance improvements via caching strategies, query optimization, and service object patterns",
    ]
)

experience(doc,
    "Ruby on Rails Developer",
    "Lightprint.com & Various Projects",
    "2008 - 2018", "USA / Pakistan",
    [
        "Delivered multiple Rails applications across healthcare, printing, food delivery, education, and SaaS domains",
        "Built Videochatapro.com (WebRTC video chat), Bringmethat.com (food delivery), Ubergrad.com (education platform)",
        "Developed Pressedgarments.com (dry cleaning SaaS), Visitdays.com (campus tours), Seedlr.com (event marketplace)",
    ]
)

# ══════════════════════════════════════════
# KEY PROJECTS
# ══════════════════════════════════════════

heading(doc, "Key Projects")

project(doc,
    "Valus.io",
    "Large-scale eCommerce grocery platform serving millions. Solidus/Spree commerce engine, real-time order tracking via ActionCable, Redis caching, AWS deployments.",
    "Rails 7, Solidus, Redis, Sidekiq, AWS"
)

project(doc,
    "Tractive",
    "GPS pet tracking system processing millions of location events. Low-latency APIs, RabbitMQ event-driven workflows, PostgreSQL optimization.",
    "Rails, RabbitMQ, PostgreSQL, Redis"
)

project(doc,
    "Doc.gr",
    "Healthcare appointment scheduling and doctor-patient management. HIPAA-compliant workflows, FHIR/HL7 data exchange, secure audit trails.",
    "Rails, React.js, FHIR/HL7, Sidekiq"
)

project(doc,
    "SpaceOS",
    "Intelligent workplace automation with IoT device integration. Multi-tenant SaaS architecture, real-time notifications via ActionCable and Turbo Streams.",
    "Rails, IoT, ActionCable, Turbo"
)

project(doc,
    "Borderhaul",
    "Cross-border logistics freight platform. Stripe Connect automated payouts, carrier matching engine, real-time shipment tracking with webhooks.",
    "Rails, Stripe Connect, ActionCable"
)

project(doc,
    "Seedlr.com",
    "Multi-tenant SaaS platform for digitizing real-life experiences into workflows. RBAC, custom feature toggles, event marketplace with payments.",
    "Rails, RBAC, SaaS, Multi-tenant"
)

project(doc,
    "CAST4k Reseller Panel",
    "Backend reseller management panel for the CAST4k IPTV platform. Stripe payment integration, subscription management, reseller onboarding, and license key generation.",
    "NestJS, TypeScript, Stripe, TypeORM, PostgreSQL"
)

project(doc,
    "FinLedger",
    "Personal finance tracking application with budget management, transaction categorization, monthly reports, and expense analytics with JWT authentication.",
    "NestJS, TypeScript, JWT, REST API"
)

project(doc,
    "Videochatapro.com",
    "Real-time video chat platform built with WebRTC and Twilio. Peer-to-peer video calls, screen sharing, chat rooms, and recording capabilities.",
    "Rails, WebRTC, Twilio, ActionCable"
)

project(doc,
    "Bringmethat.com",
    "On-demand food delivery marketplace connecting restaurants with customers. Order management, driver tracking, and Stripe payment processing.",
    "Rails, Stripe, Elasticsearch, Redis"
)

# ══════════════════════════════════════════
# EDUCATION & LANGUAGES
# ══════════════════════════════════════════

heading(doc, "Education")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(1)
run = p.add_run("Master's in Software Project Management")
run.font.size = Pt(10.5)
run.font.bold = True
run.font.color.rgb = DARK
run.font.name = 'Calibri'
run = p.add_run("  -  ")
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
run = p.add_run("FAST National University (NUCES)")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = BLUE
run.font.name = 'Calibri'
run = p.add_run("  |  2012 - 2014")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(2)
p.paragraph_format.space_after = Pt(1)
run = p.add_run("Bachelor of CS & Engineering (BCSE)")
run.font.size = Pt(10.5)
run.font.bold = True
run.font.color.rgb = DARK
run.font.name = 'Calibri'
run = p.add_run("  -  ")
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
run = p.add_run("University of Engineering and Technology (UET)")
run.font.size = Pt(10)
run.font.bold = True
run.font.color.rgb = BLUE
run.font.name = 'Calibri'
run = p.add_run("  |  2002 - 2006")
run.font.size = Pt(9)
run.font.color.rgb = GRAY
run.font.name = 'Calibri'

# Languages on same line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after = Pt(1)
run = p.add_run("Languages:  ")
run.font.size = Pt(9.5)
run.font.bold = True
run.font.color.rgb = DARK
run.font.name = 'Calibri'
run = p.add_run("English (Proficient)  |  Urdu (Native)")
run.font.size = Pt(9.5)
run.font.color.rgb = TEXT
run.font.name = 'Calibri'


# ══════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════

output_path = os.path.join(os.path.dirname(__file__), "Hamid Raza - CV.docx")
doc.save(output_path)
print(f"CV saved to: {output_path}")
